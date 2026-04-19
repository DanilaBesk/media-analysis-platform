from pathlib import Path

from docx import Document

from telegram_transcriber_bot.documents import (
    _is_human_readable_title,
    _looks_like_machine_file_name,
    build_transcript_markdown,
    build_source_label,
    build_transcript_title,
    format_timestamp,
    normalize_report_markdown,
    write_report_docx,
    write_transcript_docx,
)
from telegram_transcriber_bot.domain import TranscriptResult, TranscriptSegment


def test_build_transcript_markdown_includes_timestamps_and_speakers() -> None:
    transcript = TranscriptResult(
        title="Demo call",
        source_label="YouTube: demo",
        segments=[
            TranscriptSegment(start_seconds=0.0, end_seconds=5.0, text="Intro", speaker="Speaker 1"),
            TranscriptSegment(start_seconds=5.0, end_seconds=12.0, text="Plan", speaker=None),
        ],
        language="ru",
        raw_text="Intro\nPlan",
    )

    markdown = build_transcript_markdown(transcript)

    assert "# Demo call" in markdown
    assert "- Источник: YouTube: demo" in markdown
    assert "- Язык: ru" in markdown
    assert "## Сегменты" in markdown
    assert "## Полный текст" in markdown
    assert "[00:00 - 00:05] Speaker 1: Intro" in markdown
    assert "[00:05 - 00:12] Фрагмент: Plan" in markdown


def test_write_transcript_docx_creates_document(tmp_path: Path) -> None:
    transcript = TranscriptResult(
        title="Weekly Sync",
        source_label="Audio: sync.mp3",
        segments=[TranscriptSegment(start_seconds=0.0, end_seconds=3.5, text="Hello team", speaker="Speaker 1")],
        language="ru",
        raw_text="Hello team",
    )

    output_path = tmp_path / "transcript.docx"
    write_transcript_docx(output_path, transcript)

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Weekly Sync" in text
    assert "Источник: sync.mp3" in text
    assert "Язык: ru" in text
    assert "Сегменты" in text
    assert "Полный текст" in text
    assert "Speaker 1" in text
    assert "Hello team" in text


def test_write_report_docx_creates_document(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_report_docx(output_path, "# Report\n\n## Topics\n\n### Detail\n\n- Theme A\n\nPlain text line")

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Report" in text
    assert "Theme A" in text
    assert "Detail" in text
    assert "Plain text line" in text


def test_write_report_docx_strips_boilerplate_and_inline_markdown(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_report_docx(
        output_path,
        (
            "Вот исследовательский отчёт на основе транскрипта:\n\n---\n\n"
            "# Исследовательский отчёт\n\n## Ключевые вопросы\n\n"
            "1. **Является ли эволюция предсказуемой?**\n\n"
            "- **Тема A**\n"
        ),
    )

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Вот исследовательский отчёт" not in text
    assert "---" not in text
    assert "**" not in text
    assert "Является ли эволюция предсказуемой?" in text
    assert "Тема A" in text


def test_normalize_report_markdown_removes_intro_and_inserts_title() -> None:
    markdown = normalize_report_markdown(
        "Вот исследовательский отчёт на основе транскрипта:\n\n---\n\n## Ключевые вопросы\n\n- Theme A\n"
    )

    assert markdown.startswith("# Исследовательский отчёт\n\n## Ключевые вопросы")
    assert "Вот исследовательский отчёт" not in markdown
    assert "---" not in markdown


def test_format_timestamp_supports_hour_format() -> None:
    assert format_timestamp(3661) == "01:01:01"


def test_format_timestamp_clamps_negative_values() -> None:
    assert format_timestamp(-5) == "00:00"


def test_machine_like_audio_name_is_hidden_from_title_and_source() -> None:
    transcript = TranscriptResult(
        title="AgADqZkAAk-zkEo.ogg",
        source_label="Audio: AgADqZkAAk-zkEo.ogg",
        segments=[],
        language="ru",
        raw_text="",
    )

    assert build_transcript_title(transcript) == "Транскрибация аудио"
    assert build_source_label(transcript) == "Аудиофайл из Telegram"


def test_human_readable_media_name_is_preserved() -> None:
    transcript = TranscriptResult(
        title="meeting_notes.ogg",
        source_label="Audio: meeting_notes.ogg",
        segments=[],
        language="ru",
        raw_text="",
    )

    assert build_transcript_title(transcript) == "meeting_notes.ogg"
    assert build_source_label(transcript) == "meeting_notes.ogg"


def test_build_transcript_title_fallbacks_cover_supported_source_types() -> None:
    youtube = TranscriptResult(
        title="",
        source_label="YouTube: https://youtu.be/demo",
        segments=[],
        language="ru",
        raw_text="",
    )
    video = TranscriptResult(
        title="AgAB12345-Qwert",
        source_label="Video: AgAB12345-Qwert.mp4",
        segments=[],
        language="ru",
        raw_text="",
    )
    unknown = TranscriptResult(
        title="",
        source_label="",
        segments=[],
        language="ru",
        raw_text="",
    )

    assert build_transcript_title(youtube) == "Транскрибация YouTube-видео"
    assert build_transcript_title(video) == "Транскрибация видео"
    assert build_transcript_title(unknown) == "Транскрибация"


def test_build_source_label_fallbacks_cover_empty_and_unknown_sources() -> None:
    youtube = TranscriptResult(title="", source_label="YouTube:", segments=[], language="ru", raw_text="")
    audio = TranscriptResult(title="", source_label="Audio:", segments=[], language="ru", raw_text="")
    video = TranscriptResult(title="", source_label="Video:", segments=[], language="ru", raw_text="")
    unknown = TranscriptResult(title="", source_label="", segments=[], language="ru", raw_text="")

    assert build_source_label(youtube) == "YouTube"
    assert build_source_label(audio) == "Аудиофайл из Telegram"
    assert build_source_label(video) == "Видеофайл из Telegram"
    assert build_source_label(unknown) == "Неизвестный источник"


def test_machine_name_heuristics_stay_strict() -> None:
    assert _looks_like_machine_file_name("AgADqZkAAk-zkEo.ogg") is True
    assert _looks_like_machine_file_name("meeting_notes.ogg") is False
    assert _looks_like_machine_file_name("plain text.mp3") is False
    assert _is_human_readable_title("meeting_notes.ogg") is True
    assert _is_human_readable_title("") is False
