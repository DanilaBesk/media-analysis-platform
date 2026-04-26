# FILE: workers/transcription/src/transcriber_worker_transcription.py
# VERSION: 1.0.0
# START_MODULE_CONTRACT
# PURPOSE: Execute claimed transcription jobs through the shared control-plane client while preserving the current transcript artifact contract.
# SCOPE: Worker claim/run orchestration, ordered-input materialization, combined-media concatenation, transcript artifact persistence, cancellation checks, and packet-local helper functions for local transcription materialization.
# DEPENDS: M-WORKER-TRANSCRIPTION, M-WORKER-COMMON, M-CONTRACTS
# LINKS: M-WORKER-TRANSCRIPTION, V-M-WORKER-TRANSCRIPTION
# ROLE: RUNTIME
# MAP_MODE: EXPORTS
# END_MODULE_CONTRACT
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v1.0.0 - Introduced the transcription worker shell and extracted the local transcript orchestration path into one packet-scoped module.
# END_CHANGE_SUMMARY
#
# START_MODULE_MAP
#   SourceObjectStore - Defines the source-download boundary for claimed worker inputs.
#   TranscriptionWorkerResult - Returns the successful worker execution evidence used by packet-local tests.
#   WorkerCancellationRequested - Signals authoritative cancellation observed by the dedicated worker loop.
#   materialize_local_source - Copies a local source into a workspace without changing current bot semantics.
#   process_local_transcription - Executes the preserved local transcription pipeline and writes plain, markdown, and DOCX artifacts.
#   runTranscriptionAggregate - Claims a batch aggregate job, merges child transcript artifacts, registers root artifacts, and finalizes.
#   runTranscription - Claims a job, executes the worker pipeline, registers artifacts, and finalizes through the shared control-plane client.
# END_MODULE_MAP

from __future__ import annotations

import json
import logging
import shutil
import subprocess
from collections.abc import Mapping
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Protocol
from urllib import request

from docx import Document

from transcriber_workers_common.api import ArtifactSummary, ClaimedJobExecution, JobApiClient, JobSnapshot, OrderedWorkerInput
from transcriber_workers_common.artifacts import ArtifactDescriptor, ArtifactObjectStore, ArtifactWriter
from transcriber_workers_common.documents import build_transcript_markdown, write_transcript_docx
from transcriber_workers_common.transcribers import _download_youtube_audio
from transcriber_workers_common.domain import SourceCandidate, TranscriptArtifacts, TranscriptResult


_LOGGER = logging.getLogger(__name__)
_LOG_MARKER_EXECUTE_TRANSCRIPTION_PIPELINE = "[WorkerTranscription][runTranscription][BLOCK_EXECUTE_TRANSCRIPTION_PIPELINE]"
_AUDIO_SUFFIXES = frozenset({".aac", ".amr", ".flac", ".m4a", ".mp3", ".ogg", ".oga", ".opus", ".wav", ".wma"})
_VIDEO_SUFFIXES = frozenset({".avi", ".m4v", ".mkv", ".mov", ".mp4", ".mpeg", ".mpg", ".webm"})

__all__ = [
    "SourceObjectStore",
    "TranscriptionWorkerResult",
    "WorkerCancellationRequested",
    "materialize_local_source",
    "process_local_transcription",
    "runTranscriptionAggregate",
    "runTranscription",
]


class SourceObjectStore(Protocol):
    def fetch_file(self, *, object_key: str, destination: Path) -> None: ...


@dataclass(frozen=True, slots=True)
class TranscriptionWorkerResult:
    execution: ClaimedJobExecution
    source: SourceCandidate
    transcript: TranscriptResult
    artifacts: TranscriptArtifacts
    artifact_descriptors: tuple[ArtifactDescriptor, ...]


@dataclass(frozen=True, slots=True)
class AggregateTranscriptSection:
    source_label: str
    child_job_id: str
    plain_text: str
    markdown_text: str


@dataclass(frozen=True, slots=True)
class AggregateTranscriptionWorkerResult:
    execution: ClaimedJobExecution
    sections: tuple[AggregateTranscriptSection, ...]
    artifacts: TranscriptArtifacts
    artifact_descriptors: tuple[ArtifactDescriptor, ...]
    diagnostics: dict[str, object]


class WorkerCancellationRequested(RuntimeError):
    pass


class SourceMaterializationError(RuntimeError):
    pass


class AggregateTranscriptionError(RuntimeError):
    pass


class MissingChildTranscriptArtifactError(AggregateTranscriptionError):
    pass


def materialize_local_source(source: SourceCandidate, workspace_dir: Path) -> SourceCandidate:
    # START_BLOCK_BLOCK_MATERIALIZE_LOCAL_SOURCE
    if not source.local_path:
        return source

    workspace_dir.mkdir(parents=True, exist_ok=True)
    if source.local_path.resolve().is_relative_to(workspace_dir.resolve()):
        return source
    destination = workspace_dir / f"source{source.local_path.suffix or '.bin'}"
    if source.local_path.resolve() != destination.resolve():
        shutil.copy2(source.local_path, destination)
    return replace(source, local_path=destination)
    # END_BLOCK_BLOCK_MATERIALIZE_LOCAL_SOURCE


def process_local_transcription(
    source: SourceCandidate,
    *,
    workspace_dir: Path,
    transcriber,
) -> tuple[SourceCandidate, TranscriptResult, TranscriptArtifacts]:
    # START_BLOCK_BLOCK_PROCESS_LOCAL_TRANSCRIPTION
    materialized_source = materialize_local_source(source, workspace_dir)
    transcript_result = transcriber.transcribe(materialized_source, workspace_dir)
    artifacts = _write_transcript_artifacts(workspace_dir, transcript_result)
    return materialized_source, transcript_result, artifacts
    # END_BLOCK_BLOCK_PROCESS_LOCAL_TRANSCRIPTION


def runTranscription(
    job_id: str,
    *,
    workspace_root: Path,
    api_client: JobApiClient,
    source_store: SourceObjectStore,
    artifact_store: ArtifactObjectStore,
    transcriber,
) -> TranscriptionWorkerResult:
    execution = api_client.claim_job(job_id, worker_kind="transcription", task_type="transcription.run")
    workspace_dir = Path(workspace_root) / execution.job_id
    workspace_dir.mkdir(parents=True, exist_ok=True)

    try:
        # START_BLOCK_BLOCK_EXECUTE_TRANSCRIPTION_PIPELINE
        _check_cancellation(api_client, execution)
        api_client.publish_progress(
            execution.job_id,
            execution_id=execution.execution_id,
            progress_stage="materializing_sources",
            progress_message="Resolving claimed transcription inputs",
        )
        source = _materialize_execution_source(execution, workspace_dir, source_store)

        _check_cancellation(api_client, execution)
        api_client.publish_progress(
            execution.job_id,
            execution_id=execution.execution_id,
            progress_stage="transcribing",
            progress_message="Running transcription pipeline",
        )
        _LOGGER.info(
            "%s job_id=%s execution_id=%s ordered_input_count=%s",
            _LOG_MARKER_EXECUTE_TRANSCRIPTION_PIPELINE,
            execution.job_id,
            execution.execution_id,
            len(execution.ordered_inputs),
        )
        materialized_source, transcript_result, artifacts = process_local_transcription(
            source,
            workspace_dir=workspace_dir,
            transcriber=transcriber,
        )

        _check_cancellation(api_client, execution)
        api_client.publish_progress(
            execution.job_id,
            execution_id=execution.execution_id,
            progress_stage="persisting_artifacts",
            progress_message="Uploading transcript artifacts",
        )
        artifact_descriptors = _persist_transcript_artifacts(execution.job_id, artifacts, artifact_store)
        _assert_required_artifacts_exist(artifacts)
        api_client.register_artifacts(
            execution.job_id,
            execution_id=execution.execution_id,
            artifacts=artifact_descriptors,
        )

        _check_cancellation(api_client, execution)
        api_client.finalize_job(
            execution.job_id,
            execution_id=execution.execution_id,
            outcome="succeeded",
            progress_stage="completed",
            progress_message="Transcript ready",
            error_code=None,
            error_message=None,
        )
        return TranscriptionWorkerResult(
            execution=execution,
            source=materialized_source,
            transcript=transcript_result,
            artifacts=artifacts,
            artifact_descriptors=artifact_descriptors,
        )
        # END_BLOCK_BLOCK_EXECUTE_TRANSCRIPTION_PIPELINE
    except WorkerCancellationRequested:
        api_client.finalize_job(
            execution.job_id,
            execution_id=execution.execution_id,
            outcome="canceled",
            progress_stage="canceled",
            progress_message="Cancellation requested",
            error_code=None,
            error_message=None,
        )
        raise
    except Exception as exc:
        api_client.finalize_job(
            execution.job_id,
            execution_id=execution.execution_id,
            outcome="failed",
            progress_stage="failed",
            progress_message="Transcription failed",
            error_code=_classify_error_code(exc),
            error_message=str(exc),
        )
        raise


def runTranscriptionAggregate(
    job_id: str,
    *,
    workspace_root: Path,
    api_client: JobApiClient,
    artifact_store: ArtifactObjectStore,
) -> AggregateTranscriptionWorkerResult:
    execution = api_client.claim_job(job_id, worker_kind="transcription", task_type="transcription.aggregate")
    workspace_dir = Path(workspace_root) / execution.job_id
    workspace_dir.mkdir(parents=True, exist_ok=True)

    try:
        _check_cancellation(api_client, execution)
        api_client.publish_progress(
            execution.job_id,
            execution_id=execution.execution_id,
            progress_stage="resolving_child_artifacts",
            progress_message="Resolving batch child transcript artifacts",
        )
        root_snapshot = api_client.get_job_snapshot(execution.job_id)
        sections, diagnostics = _load_aggregate_sections(
            execution,
            root_snapshot=root_snapshot,
            workspace_dir=workspace_dir,
            api_client=api_client,
        )

        _check_cancellation(api_client, execution)
        api_client.publish_progress(
            execution.job_id,
            execution_id=execution.execution_id,
            progress_stage="merging_transcripts",
            progress_message="Merging batch transcripts",
        )
        artifacts = _write_aggregate_transcript_artifacts(workspace_dir, sections)

        _check_cancellation(api_client, execution)
        api_client.publish_progress(
            execution.job_id,
            execution_id=execution.execution_id,
            progress_stage="persisting_artifacts",
            progress_message="Uploading aggregate transcript artifacts",
        )
        artifact_descriptors = (
            *_persist_transcript_artifacts(execution.job_id, artifacts, artifact_store),
            *_persist_aggregate_metadata_artifacts(execution, sections, diagnostics, artifact_store),
        )
        _assert_required_artifacts_exist(artifacts)
        api_client.register_artifacts(
            execution.job_id,
            execution_id=execution.execution_id,
            artifacts=artifact_descriptors,
        )

        _check_cancellation(api_client, execution)
        api_client.finalize_job(
            execution.job_id,
            execution_id=execution.execution_id,
            outcome="succeeded",
            progress_stage="completed",
            progress_message="Aggregate transcript ready",
            error_code=None,
            error_message=None,
        )
        return AggregateTranscriptionWorkerResult(
            execution=execution,
            sections=sections,
            artifacts=artifacts,
            artifact_descriptors=artifact_descriptors,
            diagnostics=diagnostics,
        )
    except WorkerCancellationRequested:
        api_client.finalize_job(
            execution.job_id,
            execution_id=execution.execution_id,
            outcome="canceled",
            progress_stage="canceled",
            progress_message="Cancellation requested",
            error_code=None,
            error_message=None,
        )
        raise
    except Exception as exc:
        api_client.finalize_job(
            execution.job_id,
            execution_id=execution.execution_id,
            outcome="failed",
            progress_stage="failed",
            progress_message="Batch transcript aggregation failed",
            error_code=_classify_error_code(exc),
            error_message=str(exc),
        )
        raise


def _write_transcript_artifacts(workspace_dir: Path, transcript_result: TranscriptResult) -> TranscriptArtifacts:
    markdown_path = workspace_dir / "transcript.md"
    text_path = workspace_dir / "transcript.txt"
    docx_path = workspace_dir / "transcript.docx"

    markdown_path.write_text(build_transcript_markdown(transcript_result), encoding="utf-8")
    text_path.write_text(transcript_result.raw_text.strip() + "\n", encoding="utf-8")
    write_transcript_docx(docx_path, transcript_result)
    return TranscriptArtifacts(
        markdown_path=markdown_path,
        docx_path=docx_path,
        text_path=text_path,
    )


def _load_aggregate_sections(
    execution: ClaimedJobExecution,
    *,
    root_snapshot: JobSnapshot,
    workspace_dir: Path,
    api_client: JobApiClient,
) -> tuple[tuple[AggregateTranscriptSection, ...], dict[str, object]]:
    ordered_labels = _resolve_aggregate_source_labels(execution, root_snapshot)
    completion_policy = _resolve_completion_policy(execution)
    child_snapshots = tuple(
        api_client.get_job_snapshot(child.job_id)
        for child in root_snapshot.children
        if child.job_type == "transcription"
    )
    children_by_label: dict[str, list[JobSnapshot]] = {}
    for child in child_snapshots:
        label = _source_label_from_snapshot(child)
        if label:
            children_by_label.setdefault(label, []).append(child)

    sections: list[AggregateTranscriptSection] = []
    skipped_sources: list[dict[str, object]] = []
    child_artifacts: list[dict[str, object]] = []
    for label in ordered_labels:
        candidates = sorted(children_by_label.get(label, ()), key=lambda item: (item.status != "succeeded", item.job_id))
        child = next((candidate for candidate in candidates if candidate.status == "succeeded"), None)
        if child is None:
            skipped_sources.append(
                {
                    "source_label": label,
                    "reason": "child_not_succeeded",
                    "child_statuses": [candidate.status for candidate in candidates],
                    "child_job_ids": [candidate.job_id for candidate in candidates],
                }
            )
            if completion_policy == "succeed_when_all_sources_succeed":
                raise AggregateTranscriptionError(f"source {label} has no succeeded transcription child")
            continue

        plain_artifact = _artifact_by_kind(child, "transcript_plain")
        markdown_artifact = _artifact_by_kind(child, "transcript_segmented_markdown")
        if plain_artifact is None or markdown_artifact is None:
            missing = []
            if plain_artifact is None:
                missing.append("transcript_plain")
            if markdown_artifact is None:
                missing.append("transcript_segmented_markdown")
            raise MissingChildTranscriptArtifactError(
                f"source {label} child {child.job_id} is missing artifact(s): {', '.join(missing)}"
            )

        child_dir = workspace_dir / "children" / f"{len(sections):02d}-{label}"
        plain_path = _download_child_artifact(api_client, plain_artifact, child_dir / "transcript.txt")
        markdown_path = _download_child_artifact(api_client, markdown_artifact, child_dir / "transcript.md")
        child_artifacts.append(
            {
                "source_label": label,
                "child_job_id": child.job_id,
                "plain_artifact_id": plain_artifact.artifact_id,
                "markdown_artifact_id": markdown_artifact.artifact_id,
            }
        )
        sections.append(
            AggregateTranscriptSection(
                source_label=label,
                child_job_id=child.job_id,
                plain_text=plain_path.read_text(encoding="utf-8").strip(),
                markdown_text=markdown_path.read_text(encoding="utf-8").strip(),
            )
        )

    if not sections:
        raise AggregateTranscriptionError("batch transcription has no eligible successful source children")

    diagnostics: dict[str, object] = {
        "diagnostics_version": "batch-transcription.aggregate.diagnostics.v1",
        "root_job_id": execution.job_id,
        "completion_policy": completion_policy,
        "ordered_source_labels": list(ordered_labels),
        "included_count": len(sections),
        "skipped_sources": skipped_sources,
        "child_artifacts": child_artifacts,
    }
    return tuple(sections), diagnostics


def _write_aggregate_transcript_artifacts(
    workspace_dir: Path,
    sections: tuple[AggregateTranscriptSection, ...],
) -> TranscriptArtifacts:
    markdown_path = workspace_dir / "transcript.md"
    text_path = workspace_dir / "transcript.txt"
    docx_path = workspace_dir / "transcript.docx"

    text_path.write_text(_build_aggregate_plain_text(sections), encoding="utf-8")
    markdown_path.write_text(_build_aggregate_markdown(sections), encoding="utf-8")
    _write_aggregate_docx(docx_path, sections)
    return TranscriptArtifacts(
        markdown_path=markdown_path,
        docx_path=docx_path,
        text_path=text_path,
    )


def _build_aggregate_plain_text(sections: tuple[AggregateTranscriptSection, ...]) -> str:
    chunks = []
    for section in sections:
        chunks.append(f"## Транскрибация {section.source_label}\n\n{section.plain_text.strip()}")
    return "\n\n".join(chunks).strip() + "\n"


def _build_aggregate_markdown(sections: tuple[AggregateTranscriptSection, ...]) -> str:
    chunks = ["# Транскрибация", ""]
    for section in sections:
        chunks.extend(
            [
                f"## Транскрибация {section.source_label}",
                "",
                section.markdown_text.strip(),
                "",
            ]
        )
    return "\n".join(chunks).strip() + "\n"


def _write_aggregate_docx(output_path: Path, sections: tuple[AggregateTranscriptSection, ...]) -> None:
    document = Document()
    document.add_heading("Транскрибация", level=0)
    for section in sections:
        document.add_heading(f"Транскрибация {section.source_label}", level=1)
        for paragraph_text in section.plain_text.splitlines():
            text = paragraph_text.strip()
            if text:
                document.add_paragraph(text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _materialize_execution_source(
    execution: ClaimedJobExecution,
    workspace_dir: Path,
    source_store: SourceObjectStore,
) -> SourceCandidate:
    ordered_inputs = tuple(sorted(execution.ordered_inputs, key=lambda item: item.position))
    if len(ordered_inputs) == 1:
        return _materialize_single_ordered_input(
            ordered_inputs[0],
            workspace_dir,
            source_store,
            source_label=_resolve_claimed_source_label(ordered_inputs[0], execution.params),
        )
    return _materialize_combined_source(ordered_inputs, workspace_dir, source_store)


def _materialize_single_ordered_input(
    ordered_input: OrderedWorkerInput,
    workspace_dir: Path,
    source_store: SourceObjectStore,
    *,
    source_label: str | None = None,
) -> SourceCandidate:
    input_dir = workspace_dir / "inputs" / f"{ordered_input.position:02d}-{ordered_input.source_id}"
    input_dir.mkdir(parents=True, exist_ok=True)
    display_name = source_label or _resolve_display_name(ordered_input)

    if ordered_input.source_kind == "youtube_url":
        if not ordered_input.source_url:
            raise SourceMaterializationError("youtube_url input must include source_url")
        return SourceCandidate(
            source_id=ordered_input.source_id,
            kind="youtube_url",
            display_name=display_name,
            url=ordered_input.source_url,
            telegram_file_id=None,
            mime_type=None,
            file_name=ordered_input.original_filename,
            file_unique_id=None,
            local_path=None,
        )

    local_path = _download_ordered_input(ordered_input, input_dir, source_store)
    return SourceCandidate(
        source_id=ordered_input.source_id,
        kind=_guess_media_kind(ordered_input),
        display_name=display_name,
        url=None,
        telegram_file_id=None,
        mime_type=None,
        file_name=ordered_input.original_filename or local_path.name,
        file_unique_id=None,
        local_path=local_path,
    )


def _materialize_combined_source(
    ordered_inputs: tuple[OrderedWorkerInput, ...],
    workspace_dir: Path,
    source_store: SourceObjectStore,
) -> SourceCandidate:
    materialized_paths: list[Path] = []
    combined_dir = workspace_dir / "combined"
    combined_dir.mkdir(parents=True, exist_ok=True)

    for ordered_input in ordered_inputs:
        item_dir = combined_dir / f"{ordered_input.position:02d}-{ordered_input.source_id}"
        item_dir.mkdir(parents=True, exist_ok=True)
        if ordered_input.source_kind == "youtube_url":
            if not ordered_input.source_url:
                raise SourceMaterializationError("youtube_url input must include source_url")
            try:
                local_path = _download_youtube_audio(ordered_input.source_url, item_dir)
            except Exception as exc:
                raise SourceMaterializationError(str(exc)) from exc
        else:
            local_path = _download_ordered_input(ordered_input, item_dir, source_store)
        materialized_paths.append(local_path)

    output_path = combined_dir / "combined.wav"
    _concatenate_media_inputs(materialized_paths, output_path)
    return SourceCandidate(
        source_id=f"{ordered_inputs[0].source_id}-combined",
        kind="telegram_audio",
        display_name="Audio: combined-inputs.wav",
        url=None,
        telegram_file_id=None,
        mime_type="audio/wav",
        file_name="combined-inputs.wav",
        file_unique_id=None,
        local_path=output_path,
    )


def _download_ordered_input(
    ordered_input: OrderedWorkerInput,
    destination_dir: Path,
    source_store: SourceObjectStore,
) -> Path:
    if not ordered_input.object_key:
        raise SourceMaterializationError(f"{ordered_input.source_kind} input must include object_key")

    destination = destination_dir / _resolve_materialized_filename(ordered_input)
    try:
        source_store.fetch_file(object_key=ordered_input.object_key, destination=destination)
    except Exception as exc:
        raise SourceMaterializationError(str(exc)) from exc
    return destination


def _concatenate_media_inputs(input_paths: list[Path], output_path: Path) -> None:
    if len(input_paths) < 2:
        raise SourceMaterializationError("combined transcription requires at least two inputs")

    filter_inputs = "".join(f"[{index}:a]" for index in range(len(input_paths)))
    command = ["ffmpeg", "-y"]
    for path in input_paths:
        command.extend(["-i", str(path)])
    command.extend(
        [
            "-filter_complex",
            f"{filter_inputs}concat=n={len(input_paths)}:v=0:a=1[outa]",
            "-map",
            "[outa]",
            "-ac",
            "1",
            "-ar",
            "16000",
            str(output_path),
        ]
    )
    completed = subprocess.run(command, capture_output=True, text=True, check=False, timeout=3600)
    if completed.returncode != 0:
        raise SourceMaterializationError(f"ffmpeg concat failed with exit code {completed.returncode}: {completed.stderr.strip()}")


def _persist_transcript_artifacts(
    job_id: str,
    artifacts: TranscriptArtifacts,
    artifact_store: ArtifactObjectStore,
) -> tuple[ArtifactDescriptor, ...]:
    writer = ArtifactWriter(job_id=job_id, object_store=artifact_store)
    return (
        writer.write_file_artifact(
            "transcript_plain",
            artifacts.text_path,
            mime_type="text/plain; charset=utf-8",
            format="txt",
        ),
        writer.write_file_artifact(
            "transcript_segmented_markdown",
            artifacts.markdown_path,
            mime_type="text/markdown; charset=utf-8",
            format="markdown",
        ),
        writer.write_file_artifact(
            "transcript_docx",
            artifacts.docx_path,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            format="docx",
        ),
    )


def _persist_aggregate_metadata_artifacts(
    execution: ClaimedJobExecution,
    sections: tuple[AggregateTranscriptSection, ...],
    diagnostics: dict[str, object],
    artifact_store: ArtifactObjectStore,
) -> tuple[ArtifactDescriptor, ...]:
    writer = ArtifactWriter(job_id=execution.job_id, object_store=artifact_store)
    manifest = {
        "manifest_version": "batch-transcription.aggregate.v1",
        "root_job_id": execution.job_id,
        "ordered_source_labels": [section.source_label for section in sections],
        "included_sources": [
            {
                "source_label": section.source_label,
                "child_job_id": section.child_job_id,
            }
            for section in sections
        ],
    }
    return (
        writer.write_text_artifact(
            "source_manifest_json",
            "source-manifest.json",
            _json_dump(manifest),
            mime_type="application/json; charset=utf-8",
            format="json",
        ),
        writer.write_text_artifact(
            "batch_diagnostics_json",
            "batch-diagnostics.json",
            _json_dump(diagnostics),
            mime_type="application/json; charset=utf-8",
            format="json",
        ),
    )


def _assert_required_artifacts_exist(artifacts: TranscriptArtifacts) -> None:
    required_paths = (
        artifacts.text_path,
        artifacts.markdown_path,
        artifacts.docx_path,
    )
    for path in required_paths:
        if not path.exists():
            raise RuntimeError(f"required transcript artifact is missing: {path}")


def _check_cancellation(api_client: JobApiClient, execution: ClaimedJobExecution) -> None:
    cancel_state = api_client.check_cancel(execution.job_id, execution_id=execution.execution_id)
    if cancel_state.cancel_requested:
        raise WorkerCancellationRequested(f"job {execution.job_id} was canceled")


def _resolve_claimed_source_label(ordered_input: OrderedWorkerInput, params: Mapping[str, object]) -> str | None:
    if ordered_input.source_label:
        return ordered_input.source_label.strip() or None

    batch_params = params.get("batch")
    if not isinstance(batch_params, Mapping):
        return None
    source_label = batch_params.get("source_label")
    if not isinstance(source_label, str):
        return None
    return source_label.strip() or None


def _resolve_display_name(ordered_input: OrderedWorkerInput) -> str:
    if ordered_input.display_name:
        return ordered_input.display_name
    if ordered_input.original_filename:
        media_prefix = "Video" if _guess_media_kind(ordered_input) == "telegram_video" else "Audio"
        return f"{media_prefix}: {ordered_input.original_filename}"
    if ordered_input.source_url:
        return f"YouTube: {ordered_input.source_url}"
    return ordered_input.source_id


def _resolve_materialized_filename(ordered_input: OrderedWorkerInput) -> str:
    if ordered_input.original_filename:
        return ordered_input.original_filename
    if ordered_input.display_name:
        normalized = ordered_input.display_name.split(":", 1)[-1].strip().replace("/", "-")
        suffix = Path(normalized).suffix
        if suffix:
            return normalized
    suffix = ".bin"
    if ordered_input.source_kind == "youtube_url":
        suffix = ".mp3"
    return f"{ordered_input.source_id}{suffix}"


def _guess_media_kind(ordered_input: OrderedWorkerInput) -> str:
    source_name = (ordered_input.original_filename or ordered_input.display_name or "").strip()
    if source_name.startswith("Video:"):
        return "telegram_video"
    if source_name.startswith("Audio:"):
        return "telegram_audio"

    suffix = Path(source_name).suffix.casefold()
    if suffix in _VIDEO_SUFFIXES:
        return "telegram_video"
    if suffix in _AUDIO_SUFFIXES:
        return "telegram_audio"
    return "telegram_audio"


def _resolve_aggregate_source_labels(execution: ClaimedJobExecution, root_snapshot: JobSnapshot) -> tuple[str, ...]:
    batch = execution.params.get("batch")
    if isinstance(batch, Mapping):
        labels = batch.get("ordered_source_labels")
        if isinstance(labels, list):
            normalized = tuple(str(label).strip() for label in labels if str(label).strip())
            if normalized:
                return normalized

    snapshot_labels = tuple(
        item.source_label.strip()
        for item in sorted(root_snapshot.source_set_items, key=lambda source: source.position)
        if item.source_label and item.source_label.strip()
    )
    if snapshot_labels:
        return snapshot_labels

    raise AggregateTranscriptionError("aggregate job has no ordered source labels")


def _resolve_completion_policy(execution: ClaimedJobExecution) -> str:
    batch = execution.params.get("batch")
    if isinstance(batch, Mapping):
        policy = str(batch.get("completion_policy") or "").strip()
        if policy:
            return policy
    return "succeed_when_all_sources_succeed"


def _source_label_from_snapshot(snapshot: JobSnapshot) -> str | None:
    for item in sorted(snapshot.source_set_items, key=lambda source: source.position):
        if item.source_label and item.source_label.strip():
            return item.source_label.strip()
    return None


def _artifact_by_kind(snapshot: JobSnapshot, artifact_kind: str) -> ArtifactSummary | None:
    for artifact in snapshot.artifacts:
        if artifact.artifact_kind == artifact_kind:
            return artifact
    return None


def _download_child_artifact(api_client: JobApiClient, artifact: ArtifactSummary, destination: Path) -> Path:
    resolution = api_client.resolve_artifact(artifact.artifact_id)
    destination.parent.mkdir(parents=True, exist_ok=True)
    with request.urlopen(resolution.download_url, timeout=60) as response:
        destination.write_bytes(response.read())
    return destination


def _json_dump(payload: dict[str, object]) -> str:
    return json.dumps(payload, ensure_ascii=False, sort_keys=True, indent=2) + "\n"


def _classify_error_code(error: Exception) -> str:
    if isinstance(error, SourceMaterializationError):
        return "source_fetch_failed"
    if isinstance(error, MissingChildTranscriptArtifactError):
        return "missing_child_artifact"
    if isinstance(error, AggregateTranscriptionError):
        return "batch_aggregation_failed"
    return "transcription_failed"
