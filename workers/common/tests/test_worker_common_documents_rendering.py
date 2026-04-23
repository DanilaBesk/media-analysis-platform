# FILE: workers/common/tests/test_worker_common_documents_rendering.py
# VERSION: 1.0.0
# START_MODULE_CONTRACT
# PURPOSE: Verify the remaining worker-common document rendering and normalization branches needed by the packet-local quality gate.
# SCOPE: Report DOCX rendering, title and source heuristics, timestamp formatting, and markdown cleanup helpers.
# DEPENDS: M-WORKER-COMMON
# LINKS: M-WORKER-COMMON, V-M-WORKER-COMMON
# ROLE: TEST
# MAP_MODE: LOCALS
# END_MODULE_CONTRACT
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v1.0.0 - Added packet-local document rendering coverage for target-service verification.
# END_CHANGE_SUMMARY
#
# START_MODULE_MAP
#   test_write_report_docx_creates_document - Verifies report headings, lists, and plain text rendering.
#   test_write_report_docx_strips_boilerplate_and_inline_markdown - Verifies boilerplate removal and inline markdown handling.
#   test_title_and_source_fallbacks_cover_supported_source_types - Verifies title/source heuristics for target artifacts.
#   test_timestamp_and_machine_name_helpers_cover_edge_cases - Verifies timestamp and machine-name helper branches.
# END_MODULE_MAP

from __future__ import annotations

from pathlib import Path

from docx import Document

from transcriber_workers_common.documents import (
    _is_human_readable_title,
    _looks_like_machine_file_name,
    build_source_label,
    build_transcript_title,
    format_timestamp,
    normalize_report_markdown,
    write_report_docx,
)
from telegram_transcriber_bot.domain import TranscriptResult


def test_write_report_docx_creates_document(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_report_docx(output_path, "# Report\n\n## Topics\n\n### Detail\n\n- Theme A\n\n1. Theme B\n\nPlain text line")

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Report" in text
    assert "Topics" in text
    assert "Detail" in text
    assert "Theme A" in text
    assert "Theme B" in text
    assert "Plain text line" in text


def test_write_report_docx_strips_boilerplate_and_inline_markdown(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_report_docx(
        output_path,
        (
            "Вот исследовательский отчёт на основе транскрипта:\n\n---\n\n"
            "# Исследовательский отчёт\n\n## Ключевые вопросы\n\n"
            "1. **Является ли эволюция предсказуемой?**\n\n"
            "- **Тема A**\n"
        ),
    )

    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Вот исследовательский отчёт" not in text
    assert "---" not in text
    assert "**" not in text
    assert "Является ли эволюция предсказуемой?" in text
    assert "Тема A" in text


def test_title_and_source_fallbacks_cover_supported_source_types() -> None:
    youtube = TranscriptResult(title="", source_label="YouTube: https://youtu.be/demo", segments=[], language="ru", raw_text="")
    youtube_empty = TranscriptResult(title="", source_label="YouTube:", segments=[], language="ru", raw_text="")
    audio = TranscriptResult(title="", source_label="Audio:", segments=[], language="ru", raw_text="")
    video = TranscriptResult(title="AgAB12345-Qwert", source_label="Video: AgAB12345-Qwert.mp4", segments=[], language="ru", raw_text="")
    unknown = TranscriptResult(title="", source_label="", segments=[], language="ru", raw_text="")

    assert build_transcript_title(youtube) == "Транскрибация YouTube-видео"
    assert build_transcript_title(audio) == "Транскрибация аудио"
    assert build_transcript_title(video) == "Транскрибация видео"
    assert build_transcript_title(unknown) == "Транскрибация"
    assert build_source_label(youtube) == "YouTube: https://youtu.be/demo"
    assert build_source_label(youtube_empty) == "YouTube"
    assert build_source_label(audio) == "Аудиофайл из Telegram"
    assert build_source_label(video) == "Видеофайл из Telegram"
    assert build_source_label(unknown) == "Неизвестный источник"


def test_timestamp_and_machine_name_helpers_cover_edge_cases() -> None:
    transcript = TranscriptResult(
        title="AgADqZkAAk-zkEo.ogg",
        source_label="Audio: AgADqZkAAk-zkEo.ogg",
        segments=[],
        language="ru",
        raw_text="",
    )

    assert build_transcript_title(transcript) == "Транскрибация аудио"
    assert build_source_label(transcript) == "Аудиофайл из Telegram"
    assert format_timestamp(3661) == "01:01:01"
    assert format_timestamp(-5) == "00:00"
    assert _looks_like_machine_file_name("AgADqZkAAk-zkEo.ogg") is True
    assert _looks_like_machine_file_name("meeting_notes.ogg") is False
    assert _is_human_readable_title("meeting_notes.ogg") is True
    assert _is_human_readable_title("") is False
    assert normalize_report_markdown("") == "# Исследовательский отчёт\n"
