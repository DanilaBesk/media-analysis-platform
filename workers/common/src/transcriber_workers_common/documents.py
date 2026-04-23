# FILE: workers/common/src/transcriber_workers_common/documents.py
# VERSION: 1.0.0
# START_MODULE_CONTRACT
# PURPOSE: Render transcript and report artifacts into markdown and DOCX outputs that later worker packets can reuse.
# SCOPE: Transcript markdown generation, transcript DOCX rendering, report markdown normalization, and report DOCX rendering.
# DEPENDS: M-WORKER-COMMON
# LINKS: M-WORKER-COMMON, V-M-WORKER-TRANSCRIPTION
# ROLE: RUNTIME
# MAP_MODE: EXPORTS
# END_MODULE_CONTRACT
#
# START_CHANGE_SUMMARY
#   LAST_CHANGE: v1.0.0 - Extracted reusable document-rendering helpers into worker-common without changing current output behavior.
# END_CHANGE_SUMMARY
#
# START_MODULE_MAP
#   build_transcript_markdown - Render segmented transcript markdown output.
#   write_transcript_docx - Render transcript DOCX output.
#   write_report_docx - Render report DOCX output from normalized markdown.
#   format_timestamp - Format transcript timestamps.
#   build_transcript_title - Derive a stable transcript title.
#   build_source_label - Humanize the transcript source label.
#   normalize_report_markdown - Remove report boilerplate and ensure a title heading exists.
# END_MODULE_MAP

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from telegram_transcriber_bot.domain import TranscriptResult
from telegram_transcriber_bot.source_labels import humanize_source_label, is_human_readable_title, looks_like_machine_file_name

__all__ = [
    "build_source_label",
    "build_transcript_markdown",
    "build_transcript_title",
    "format_timestamp",
    "normalize_report_markdown",
    "write_report_docx",
    "write_transcript_docx",
]


# START_CONTRACT: build_transcript_markdown
# PURPOSE: Render transcript segments and raw text into the canonical markdown artifact.
# INPUTS: { transcript: TranscriptResult - Transcript data to render }
# OUTPUTS: { str - Canonical transcript markdown }
# SIDE_EFFECTS: none
# LINKS: M-WORKER-COMMON, V-M-WORKER-TRANSCRIPTION
# END_CONTRACT: build_transcript_markdown
def build_transcript_markdown(transcript: TranscriptResult) -> str:
    # START_BLOCK_BLOCK_RENDER_TRANSCRIPT_MARKDOWN
    title = build_transcript_title(transcript)
    source_label = build_source_label(transcript)
    lines = [
        f"# {title}",
        "",
        f"- Источник: {source_label}",
        f"- Язык: {transcript.language}",
        "",
        "## Сегменты",
        "",
    ]

    for segment in transcript.segments:
        speaker = segment.speaker or "Фрагмент"
        lines.append(
            f"[{format_timestamp(segment.start_seconds)} - {format_timestamp(segment.end_seconds)}] "
            f"{speaker}: {segment.text.strip()}"
        )

    lines.extend(["", "## Полный текст", "", transcript.raw_text.strip()])
    return "\n".join(lines).strip() + "\n"
    # END_BLOCK_BLOCK_RENDER_TRANSCRIPT_MARKDOWN


# START_CONTRACT: write_transcript_docx
# PURPOSE: Render transcript content into the canonical DOCX artifact.
# INPUTS: { output_path: Path - Destination DOCX path, transcript: TranscriptResult - Transcript data to render }
# OUTPUTS: { None - The DOCX file is written to disk }
# SIDE_EFFECTS: filesystem writes
# LINKS: M-WORKER-COMMON, V-M-WORKER-TRANSCRIPTION
# END_CONTRACT: write_transcript_docx
def write_transcript_docx(output_path: Path, transcript: TranscriptResult) -> None:
    document = Document()
    document.add_heading(build_transcript_title(transcript), level=0)
    document.add_paragraph(f"Источник: {build_source_label(transcript)}")
    document.add_paragraph(f"Язык: {transcript.language}")

    document.add_heading("Сегменты", level=1)
    for segment in transcript.segments:
        speaker = segment.speaker or "Фрагмент"
        document.add_paragraph(
            f"[{format_timestamp(segment.start_seconds)} - {format_timestamp(segment.end_seconds)}] "
            f"{speaker}: {segment.text.strip()}"
        )

    document.add_heading("Полный текст", level=1)
    document.add_paragraph(transcript.raw_text.strip())
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


# START_CONTRACT: write_report_docx
# PURPOSE: Render normalized report markdown into the canonical report DOCX artifact.
# INPUTS: { output_path: Path - Destination DOCX path, markdown_content: str - Source markdown content }
# OUTPUTS: { None - The DOCX file is written to disk }
# SIDE_EFFECTS: filesystem writes
# LINKS: M-WORKER-COMMON, V-M-WORKER-REPORT
# END_CONTRACT: write_report_docx
def write_report_docx(output_path: Path, markdown_content: str) -> None:
    normalized_content = normalize_report_markdown(markdown_content)
    document = Document()
    for raw_line in normalized_content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(_strip_inline_markdown(line[2:].strip()), level=0)
            continue
        if line.startswith("## "):
            document.add_heading(_strip_inline_markdown(line[3:].strip()), level=1)
            continue
        if line.startswith("### "):
            document.add_heading(_strip_inline_markdown(line[4:].strip()), level=2)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _append_inline_markdown(paragraph, line[2:].strip())
            continue
        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if ordered_match:
            paragraph = document.add_paragraph(style="List Number")
            _append_inline_markdown(paragraph, ordered_match.group(2).strip())
            continue
        paragraph = document.add_paragraph()
        _append_inline_markdown(paragraph, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


# START_CONTRACT: format_timestamp
# PURPOSE: Format transcript seconds into canonical timestamp strings.
# INPUTS: { total_seconds: float - Segment time value in seconds }
# OUTPUTS: { str - `MM:SS` or `HH:MM:SS` formatted timestamp }
# SIDE_EFFECTS: none
# LINKS: M-WORKER-COMMON
# END_CONTRACT: format_timestamp
def format_timestamp(total_seconds: float) -> str:
    rounded = max(0, int(total_seconds))
    minutes, seconds = divmod(rounded, 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    return f"{minutes:02d}:{seconds:02d}"


# START_CONTRACT: build_transcript_title
# PURPOSE: Derive the canonical transcript title from explicit or fallback source metadata.
# INPUTS: { transcript: TranscriptResult - Transcript metadata and source label }
# OUTPUTS: { str - Human-readable transcript title }
# SIDE_EFFECTS: none
# LINKS: M-WORKER-COMMON
# END_CONTRACT: build_transcript_title
def build_transcript_title(transcript: TranscriptResult) -> str:
    explicit_title = transcript.title.strip()
    if _is_human_readable_title(explicit_title):
        return explicit_title

    source_label = transcript.source_label.strip()
    if source_label.startswith("YouTube:"):
        return "Транскрибация YouTube-видео"
    if source_label.startswith("Audio:"):
        return "Транскрибация аудио"
    if source_label.startswith("Video:"):
        return "Транскрибация видео"
    return "Транскрибация"


# START_CONTRACT: build_source_label
# PURPOSE: Humanize the transcript source label for artifact output.
# INPUTS: { transcript: TranscriptResult - Transcript metadata and source label }
# OUTPUTS: { str - Humanized source label }
# SIDE_EFFECTS: none
# LINKS: M-WORKER-COMMON
# END_CONTRACT: build_source_label
def build_source_label(transcript: TranscriptResult) -> str:
    return humanize_source_label(transcript.source_label)


def _is_human_readable_title(value: str) -> bool:
    return is_human_readable_title(value)


def _looks_like_machine_file_name(value: str) -> bool:
    return looks_like_machine_file_name(value)


# START_CONTRACT: normalize_report_markdown
# PURPOSE: Remove boilerplate and ensure a stable title heading before DOCX rendering.
# INPUTS: { markdown_content: str - Raw report markdown }
# OUTPUTS: { str - Normalized report markdown }
# SIDE_EFFECTS: none
# LINKS: M-WORKER-COMMON, V-M-WORKER-REPORT
# END_CONTRACT: normalize_report_markdown
def normalize_report_markdown(markdown_content: str) -> str:
    cleaned_lines: list[str] = []
    saw_heading = False
    pending_blank = False

    for raw_line in markdown_content.splitlines():
        line = raw_line.strip()
        if not line:
            if cleaned_lines:
                pending_blank = True
            continue
        if _is_report_boilerplate(line):
            continue
        if pending_blank and cleaned_lines:
            cleaned_lines.append("")
            pending_blank = False
        if line.startswith("# "):
            saw_heading = True
        cleaned_lines.append(line)

    if not cleaned_lines:
        return "# Исследовательский отчёт\n"
    if not saw_heading:
        cleaned_lines.insert(0, "# Исследовательский отчёт")
        cleaned_lines.insert(1, "")
    return "\n".join(cleaned_lines).strip() + "\n"


def _is_report_boilerplate(line: str) -> bool:
    normalized = line.casefold()
    return normalized in {"---", "—"} or normalized.startswith("вот исследовательский отч")


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _append_inline_markdown(paragraph, text: str) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])
