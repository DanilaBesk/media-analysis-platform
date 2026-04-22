from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from telegram_transcriber_bot.domain import TranscriptResult
from telegram_transcriber_bot.source_labels import humanize_source_label, is_human_readable_title, looks_like_machine_file_name


def build_transcript_markdown(transcript: TranscriptResult) -> str:
    title = build_transcript_title(transcript)
    source_label = build_source_label(transcript)
    lines = [
        f"# {title}",
        "",
        f"- Источник: {source_label}",
        f"- Язык: {transcript.language}",
        "",
        "## Сегменты",
        "",
    ]

    for segment in transcript.segments:
        speaker = segment.speaker or "Фрагмент"
        lines.append(
            f"[{format_timestamp(segment.start_seconds)} - {format_timestamp(segment.end_seconds)}] "
            f"{speaker}: {segment.text.strip()}"
        )

    lines.extend(["", "## Полный текст", "", transcript.raw_text.strip()])
    return "\n".join(lines).strip() + "\n"


def write_transcript_docx(output_path: Path, transcript: TranscriptResult) -> None:
    document = Document()
    document.add_heading(build_transcript_title(transcript), level=0)
    document.add_paragraph(f"Источник: {build_source_label(transcript)}")
    document.add_paragraph(f"Язык: {transcript.language}")

    document.add_heading("Сегменты", level=1)
    for segment in transcript.segments:
        speaker = segment.speaker or "Фрагмент"
        document.add_paragraph(
            f"[{format_timestamp(segment.start_seconds)} - {format_timestamp(segment.end_seconds)}] "
            f"{speaker}: {segment.text.strip()}"
        )

    document.add_heading("Полный текст", level=1)
    document.add_paragraph(transcript.raw_text.strip())
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def write_report_docx(output_path: Path, markdown_content: str) -> None:
    normalized_content = normalize_report_markdown(markdown_content)
    document = Document()
    for raw_line in normalized_content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(_strip_inline_markdown(line[2:].strip()), level=0)
            continue
        if line.startswith("## "):
            document.add_heading(_strip_inline_markdown(line[3:].strip()), level=1)
            continue
        if line.startswith("### "):
            document.add_heading(_strip_inline_markdown(line[4:].strip()), level=2)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _append_inline_markdown(paragraph, line[2:].strip())
            continue
        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if ordered_match:
            paragraph = document.add_paragraph(style="List Number")
            _append_inline_markdown(paragraph, ordered_match.group(2).strip())
            continue
        paragraph = document.add_paragraph()
        _append_inline_markdown(paragraph, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def format_timestamp(total_seconds: float) -> str:
    rounded = max(0, int(total_seconds))
    minutes, seconds = divmod(rounded, 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    return f"{minutes:02d}:{seconds:02d}"


def build_transcript_title(transcript: TranscriptResult) -> str:
    explicit_title = transcript.title.strip()
    if _is_human_readable_title(explicit_title):
        return explicit_title

    source_label = transcript.source_label.strip()
    if source_label.startswith("YouTube:"):
        return "Транскрибация YouTube-видео"
    if source_label.startswith("Audio:"):
        return "Транскрибация аудио"
    if source_label.startswith("Video:"):
        return "Транскрибация видео"
    return "Транскрибация"


def build_source_label(transcript: TranscriptResult) -> str:
    return humanize_source_label(transcript.source_label)


def _is_human_readable_title(value: str) -> bool:
    return is_human_readable_title(value)


def _looks_like_machine_file_name(value: str) -> bool:
    return looks_like_machine_file_name(value)


def normalize_report_markdown(markdown_content: str) -> str:
    cleaned_lines: list[str] = []
    saw_heading = False
    pending_blank = False

    for raw_line in markdown_content.splitlines():
        line = raw_line.strip()
        if not line:
            if cleaned_lines:
                pending_blank = True
            continue
        if _is_report_boilerplate(line):
            continue
        if pending_blank and cleaned_lines:
            cleaned_lines.append("")
            pending_blank = False
        if line.startswith("# "):
            saw_heading = True
        cleaned_lines.append(line)

    if not cleaned_lines:
        return "# Исследовательский отчёт\n"
    if not saw_heading:
        cleaned_lines.insert(0, "# Исследовательский отчёт")
        cleaned_lines.insert(1, "")
    return "\n".join(cleaned_lines).strip() + "\n"


def _is_report_boilerplate(line: str) -> bool:
    normalized = line.casefold()
    return normalized in {"---", "—"} or normalized.startswith("вот исследовательский отч")


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _append_inline_markdown(paragraph, text: str) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])
